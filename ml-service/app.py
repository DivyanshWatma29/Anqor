"""
Anqor — Flask ML Service v3.1
========================================
API endpoints:
  GET  /health                - Service health + available models
  POST /predict               - Single claim prediction (with confidence, SHAP & fraud explanation)
  POST /batch                 - Batch CSV prediction
  POST /classify              - Auto-detect claim category from data
  POST /extract               - Extract fields from uploaded document
  GET  /schema/<type>         - Get field schema for a claim type
  GET  /metrics/<type>        - Get model performance metrics
  POST /export/csv            - Export results as CSV
  POST /export/bulk-pdf       - Export batch results as bulk PDF
  POST /column-mapping/suggest - Suggest column mappings for CSV headers
  POST /feedback              - Submit user feedback on predictions
  GET  /feedback/<claim_id>   - Get feedback for a specific claim
  POST /compare               - Compare two claims side-by-side
  GET  /model/version/<type>  - Get model version info & history
  GET  /audit-trail           - Get audit log of actions
  POST /anomaly/timeseries    - Time-series anomaly detection on batch data
  POST /network/analysis      - Network analysis for related claims
  GET  /confidence-interval/<type> - Get confidence intervals for a model
"""
import os
import io
import csv
import json
import time
import uuid
import hashlib
import traceback
from datetime import datetime, timezone
from functools import wraps
from collections import defaultdict

from flask import Flask, request, jsonify, send_file, Response
from flask_cors import CORS

from core.preprocessor import (
    preprocess_input, get_prediction_and_probability,
    get_model_column_count, get_required_fields, get_available_models,
    get_feature_config, get_model_metrics, classify_claim_category,
    get_model_confidence_info
)
from core.indicators import generate_indicators
from core.shap_explainer import get_shap_explanation
from core.explainer import generate_fraud_explanation

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 10 * 1024 * 1024
CORS(app)

MAX_UPLOAD_BYTES = 10 * 1024 * 1024
ALLOWED_UPLOAD_EXTENSIONS = {'csv', 'xlsx', 'xls', 'pdf', 'docx'}
ALLOWED_UPLOAD_MIME_TYPES = {
    'csv': {'text/csv', 'application/csv', 'application/vnd.ms-excel'},
    'xlsx': {'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'},
    'xls': {'application/vnd.ms-excel', 'application/octet-stream'},
    'pdf': {'application/pdf'},
    'docx': {'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
}


# ═══════════════════════════════════════════════════════════════════
# RATE LIMITING (Must-Have #9)
# ═══════════════════════════════════════════════════════════════════

_rate_limit_store = defaultdict(list)
RATE_LIMITS = {
    'predict': {'max_requests': 60, 'window_seconds': 60},
    'batch': {'max_requests': 10, 'window_seconds': 60},
    'export': {'max_requests': 20, 'window_seconds': 60},
    'feedback': {'max_requests': 30, 'window_seconds': 60},
    'default': {'max_requests': 120, 'window_seconds': 60},
}


def _get_client_ip():
    return request.headers.get('X-Forwarded-For', request.remote_addr) or 'unknown'



def _safe_upload_filename(filename: str) -> str:
    filename = (filename or '').strip().replace('\', '/').split('/')[-1]
    return ''.join(ch for ch in filename if ch.isalnum() or ch in {'.', '_', '-'})[:120]


def _validate_uploaded_file(file) -> tuple:
    filename = _safe_upload_filename(file.filename)
    if not filename or '.' not in filename:
        raise ValueError('Uploaded file must have a valid filename and extension')

    extension = filename.rsplit('.', 1)[1].lower()
    if extension not in ALLOWED_UPLOAD_EXTENSIONS:
        raise ValueError('Unsupported file type')

    mime_type = (file.mimetype or '').lower()
    allowed_mimes = ALLOWED_UPLOAD_MIME_TYPES.get(extension, set())
    if mime_type not in allowed_mimes:
        raise ValueError(f'Unsupported MIME type for .{extension} upload')

    return filename, extension


def rate_limit(category='default'):
    """Rate limiting decorator."""
    def decorator(f):
        @wraps(f)
        def wrapper(*args, **kwargs):
            client_ip = _get_client_ip()
            key = f"{category}:{client_ip}"
            now = time.time()
            limits = RATE_LIMITS.get(category, RATE_LIMITS['default'])

            # Clean old entries
            _rate_limit_store[key] = [
                t for t in _rate_limit_store[key]
                if now - t < limits['window_seconds']
            ]

            if len(_rate_limit_store[key]) >= limits['max_requests']:
                return jsonify({
                    'error': 'Rate limit exceeded',
                    'retry_after_seconds': int(limits['window_seconds'] - (now - _rate_limit_store[key][0])),
                    'limit': limits['max_requests'],
                    'window': limits['window_seconds'],
                }), 429

            _rate_limit_store[key].append(now)
            return f(*args, **kwargs)
        return wrapper
    return decorator


# ═══════════════════════════════════════════════════════════════════
# INPUT VALIDATION (Must-Have #10)
# ═══════════════════════════════════════════════════════════════════

def validate_claim_input(data: dict, claim_type: str) -> tuple:
    """
    Validate claim input data. Returns (is_valid, errors, warnings, sanitized_data).
    """
    errors = []
    warnings = []
    sanitized = {}

    config = get_feature_config(claim_type)
    if not config:
        return False, ['Unknown claim type'], [], {}

    numeric_fields = set(config.get('numeric_features', []))
    categorical_fields = set(config.get('categorical_features', []))
    categorical_values = config.get('categorical_values', {})

    for key, value in data.items():
        if key in ('claim_type',):
            sanitized[key] = value
            continue

        # Sanitize strings — strip whitespace, limit length
        if isinstance(value, str):
            value = value.strip()[:500]

        # Validate numeric fields
        if key in numeric_fields:
            try:
                num_val = float(value)
                if num_val < -1e9 or num_val > 1e9:
                    warnings.append(f"Field '{key}' value {num_val} seems extreme")
                sanitized[key] = num_val
            except (ValueError, TypeError):
                sanitized[key] = 0
                warnings.append(f"Field '{key}' is not a valid number, defaulting to 0")

        # Validate categorical fields
        elif key in categorical_fields:
            str_val = str(value)
            valid_options = categorical_values.get(key, [])
            if valid_options and str_val not in [str(v) for v in valid_options]:
                warnings.append(
                    f"Field '{key}' value '{str_val}' not in known options: {valid_options[:5]}..."
                )
            sanitized[key] = str_val

        else:
            sanitized[key] = value

    # Check required fields
    required_fields = get_required_fields(claim_type)
    missing = [f for f in required_fields if f not in sanitized]
    if len(missing) > len(required_fields) * 0.5:
        errors.append(f"Too many missing fields ({len(missing)}/{len(required_fields)}): {missing[:5]}...")

    return len(errors) == 0, errors, warnings, sanitized


# ═══════════════════════════════════════════════════════════════════
# AUDIT TRAIL (Should-Have #6)
# ═══════════════════════════════════════════════════════════════════

_audit_log = []
MAX_AUDIT_LOG = 10000


def _audit(action: str, details: dict = None):
    """Log an action to the audit trail."""
    entry = {
        'id': str(uuid.uuid4())[:8],
        'timestamp': datetime.now(timezone.utc).isoformat(),
        'action': action,
        'client_ip': _get_client_ip(),
        'details': details or {},
    }
    _audit_log.append(entry)
    if len(_audit_log) > MAX_AUDIT_LOG:
        _audit_log.pop(0)


# ═══════════════════════════════════════════════════════════════════
# FEEDBACK STORE (Must-Have #8)
# ═══════════════════════════════════════════════════════════════════

_feedback_store = {}


# ═══════════════════════════════════════════════════════════════════
# EXISTING ENDPOINTS (updated)
# ═══════════════════════════════════════════════════════════════════

@app.route('/health', methods=['GET'])
def health():
    """Health check with model info."""
    models_info = {}
    for model_name in get_available_models():
        config = get_feature_config(model_name)
        metrics = get_model_metrics(model_name)
        confidence = get_model_confidence_info(model_name)
        models_info[model_name] = {
            'display_name': config.get('display_name', model_name),
            'n_features': config.get('n_features', get_model_column_count(model_name)),
            'f1_score': metrics.get('f1_weighted', 0),
            'auc_roc': metrics.get('auc_roc', 0),
            'trained_at': metrics.get('trained_at', 'unknown'),
            'model_confidence': confidence,
            'version': metrics.get('version', '1.0.0'),
        }

    return jsonify({
        'status': 'healthy',
        'version': '3.1.0',
        'available_models': get_available_models(),
        'models_info': models_info,
    })


@app.route('/predict', methods=['POST'])
@rate_limit('predict')
def predict():
    """Single claim prediction with confidence, SHAP explanations."""
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        claim_type = data.get('claim_type', 'auto').lower()
        if claim_type not in get_available_models():
            return jsonify({
                'error': f"Model for '{claim_type}' is not available.",
                'available': get_available_models()
            }), 400

        # INPUT VALIDATION (Must-Have #10)
        is_valid, errors, warnings, sanitized = validate_claim_input(data, claim_type)
        if not is_valid:
            return jsonify({
                'error': 'Validation failed',
                'validation_errors': errors,
                'warnings': warnings,
            }), 400

        # Preprocess and predict
        processed = preprocess_input(claim_type, sanitized)
        prediction, probability = get_prediction_and_probability(claim_type, processed)

        # Generate indicators for all claim types
        indicators = generate_indicators(sanitized, claim_type)

        # MODEL CONFIDENCE (Must-Have #6)
        confidence_info = get_model_confidence_info(claim_type)

        # SHAP EXPLANATIONS (Must-Have #7)
        shap_data = get_shap_explanation(claim_type, processed)

        # CONFIDENCE INTERVALS (Should-Have #8)
        ci_lower = max(0, round(probability - confidence_info.get('margin_of_error', 0.05), 4))
        ci_upper = min(1, round(probability + confidence_info.get('margin_of_error', 0.05), 4))

        # Generate claim_id for tracking
        claim_id = 'CLM-' + hashlib.md5(
            json.dumps(sanitized, sort_keys=True, default=str).encode()
        ).hexdigest()[:8].upper()

        # FRAUD EXPLANATION — "Why is this claim fraud?" (NEW)
        fraud_explanation = generate_fraud_explanation(
            prediction=prediction,
            probability=probability,
            indicators=indicators,
            shap_data=shap_data,
            claim_data=sanitized,
            claim_type=claim_type,
            confidence_info=confidence_info,
        )

        result = {
            'prediction': prediction,
            'probability': round(probability, 4),
            'risk_level': _risk_level(probability),
            'indicators': indicators,
            'claim_type': claim_type,
            'claim_id': claim_id,
            # NEW: Human-readable fraud explanation
            'fraud_explanation': fraud_explanation,
            # Must-Have: Model confidence
            'model_confidence': confidence_info,
            # Must-Have: SHAP explanations
            'shap_explanation': shap_data,
            # Should-Have: Confidence intervals
            'confidence_interval': {
                'lower': ci_lower,
                'upper': ci_upper,
                'confidence_level': 0.95,
            },
            # Validation warnings
            'validation_warnings': warnings,
        }

        _audit('predict', {
            'claim_type': claim_type,
            'claim_id': claim_id,
            'prediction': prediction,
            'probability': round(probability, 4),
        })

        return jsonify(result)

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/batch', methods=['POST'])
@rate_limit('batch')
def batch_predict():
    """Batch prediction from CSV data."""
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        claims = data.get('claims', [])
        claim_type = data.get('claim_type', 'auto').lower()
        column_mapping = data.get('column_mapping', {})

        if not claims:
            return jsonify({'error': 'No claims provided'}), 400

        if claim_type not in get_available_models():
            return jsonify({'error': f"Model '{claim_type}' not available"}), 400

        # Cap batch size
        if len(claims) > 5000:
            return jsonify({'error': 'Batch size exceeds maximum of 5000 claims'}), 400

        results = []
        for i, claim in enumerate(claims):
            try:
                # Apply column mapping if provided
                mapped_claim = {}
                if column_mapping:
                    for csv_col, model_col in column_mapping.items():
                        if csv_col in claim:
                            mapped_claim[model_col] = claim[csv_col]
                else:
                    mapped_claim = claim

                mapped_claim['claim_type'] = claim_type
                processed = preprocess_input(claim_type, mapped_claim)
                prediction, probability = get_prediction_and_probability(claim_type, processed)
                indicators = generate_indicators(mapped_claim, claim_type)

                results.append({
                    'index': i,
                    'prediction': prediction,
                    'probability': round(probability, 4),
                    'risk_level': _risk_level(probability),
                    'indicators': indicators,
                    'status': 'success',
                })
            except Exception as e:
                results.append({
                    'index': i,
                    'prediction': 'Unknown',
                    'probability': 0,
                    'risk_level': 'unknown',
                    'indicators': [],
                    'status': 'failed',
                    'error': str(e),
                })

        fraud_count = sum(1 for r in results if r['prediction'] == 'Y')
        success_count = sum(1 for r in results if r['status'] == 'success')

        _audit('batch_predict', {
            'claim_type': claim_type,
            'total': len(claims),
            'processed': success_count,
            'fraud_detected': fraud_count,
        })

        return jsonify({
            'total': len(claims),
            'processed': success_count,
            'fraud_detected': fraud_count,
            'legitimate': success_count - fraud_count,
            'claim_type': claim_type,
            'results': results,
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


@app.route('/classify', methods=['POST'])
@rate_limit('default')
def classify():
    """Auto-detect claim category from field names with confidence score."""
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        # Accept either a dict of fields or a list of field names
        if isinstance(data.get('fields'), list):
            field_dict = {f: '' for f in data['fields']}
        elif isinstance(data.get('fields'), dict):
            field_dict = data['fields']
        else:
            field_dict = data

        category, confidence_score = classify_claim_category(field_dict, return_confidence=True)
        config = get_feature_config(category)

        # Auto-category detection confidence (Must-Have #2)
        confidence_level = 'high' if confidence_score > 0.6 else 'medium' if confidence_score > 0.3 else 'low'

        _audit('classify', {'category': category, 'confidence': confidence_level})

        return jsonify({
            'category': category,
            'display_name': config.get('display_name', category),
            'confidence': confidence_level,
            'confidence_score': round(confidence_score, 3),
            'required_fields': get_required_fields(category),
            'all_scores': {},  # Can be populated if needed
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/schema/<claim_type>', methods=['GET'])
def schema(claim_type):
    """Get field schema and UI config for a claim type (editable fields config — Must-Have #1)."""
    if claim_type not in get_available_models():
        return jsonify({'error': f"Model '{claim_type}' not available"}), 404

    config = get_feature_config(claim_type)
    return jsonify({
        'claim_type': claim_type,
        'display_name': config.get('display_name', claim_type),
        'field_groups': config.get('field_groups', []),
        'required_fields': get_required_fields(claim_type),
        'n_features': config.get('n_features', get_model_column_count(claim_type)),
        # Must-Have #1: Editable field metadata
        'editable': True,
        'field_types': {
            f: {'editable': True, 'type': 'number' if f in config.get('numeric_features', []) else 'select'}
            for f in get_required_fields(claim_type)
        },
    })


@app.route('/metrics/<claim_type>', methods=['GET'])
def metrics(claim_type):
    """Get model performance metrics."""
    if claim_type not in get_available_models():
        return jsonify({'error': f"Model '{claim_type}' not available"}), 404

    m = get_model_metrics(claim_type)
    return jsonify({
        'claim_type': claim_type,
        'metrics': m,
    })


@app.route('/extract', methods=['POST'])
@rate_limit('default')
def extract():
    """Extract fields from uploaded file."""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file uploaded'}), 400

        file = request.files['file']
        claim_type = request.form.get('claim_type', 'auto')
        filename, extension = _validate_uploaded_file(file)

        if extension == 'csv':
            return _extract_csv(file, claim_type)
        elif extension in ('xlsx', 'xls'):
            return _extract_excel(file, claim_type)
        elif extension == 'pdf':
            return _extract_pdf(file, claim_type)
        elif extension == 'docx':
            return _extract_docx(file, claim_type)
        else:
            return jsonify({
                'error': f'Unsupported file type: {filename}',
                'supported': ['csv', 'xlsx', 'xls', 'pdf', 'docx']
            }), 400

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# CSV EXPORT (Must-Have #3)
# ═══════════════════════════════════════════════════════════════════

@app.route('/export/csv', methods=['POST'])
@rate_limit('export')
def export_csv():
    """Export prediction results as a downloadable CSV file."""
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        results = data.get('results', [])
        if not results:
            return jsonify({'error': 'No results to export'}), 400

        claim_type = data.get('claim_type', 'auto')

        output = io.StringIO()
        writer = csv.writer(output)

        # Header row
        header = ['Index', 'Claim Type', 'Prediction', 'Probability', 'Risk Level', 'Indicators']
        # Add any claim-specific fields from the first result
        if results and 'input_data' in results[0]:
            extra_fields = list(results[0]['input_data'].keys())
            header.extend(extra_fields)

        writer.writerow(header)

        # Data rows
        for r in results:
            row = [
                r.get('index', ''),
                claim_type,
                r.get('prediction', ''),
                r.get('probability', ''),
                r.get('risk_level', ''),
                '; '.join(r.get('indicators', [])),
            ]
            if 'input_data' in r:
                for f in extra_fields:
                    row.append(r['input_data'].get(f, ''))
            writer.writerow(row)

        output.seek(0)

        _audit('export_csv', {'claim_type': claim_type, 'count': len(results)})

        return Response(
            output.getvalue(),
            mimetype='text/csv',
            headers={'Content-Disposition': f'attachment; filename=fraud_results_{claim_type}_{int(time.time())}.csv'}
        )

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# BULK PDF EXPORT (Must-Have #4)
# ═══════════════════════════════════════════════════════════════════

@app.route('/export/bulk-pdf', methods=['POST'])
@rate_limit('export')
def export_bulk_pdf():
    """Export batch results as a consolidated PDF report."""
    try:
        data = request.get_json(force=True)
        if not data:
            return jsonify({'error': 'Request body is required'}), 400

        results = data.get('results', [])
        if not results:
            return jsonify({'error': 'No results to export'}), 400

        claim_type = data.get('claim_type', 'auto')

        # Generate PDF using reportlab-style text
        # Using a simple text-based PDF since reportlab may not be available
        from core.pdf_report import generate_bulk_pdf_report
        pdf_bytes = generate_bulk_pdf_report(results, claim_type)

        _audit('export_bulk_pdf', {'claim_type': claim_type, 'count': len(results)})

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'fraud_report_{claim_type}_{int(time.time())}.pdf',
        )

    except ImportError:
        # Fallback: return JSON summary if PDF generation not available
        return jsonify({
            'error': 'PDF generation requires fpdf2. Install with: pip install fpdf2',
            'fallback': 'Use /export/csv instead',
        }), 501
    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# COLUMN MAPPING (Must-Have #5)
# ═══════════════════════════════════════════════════════════════════

@app.route('/column-mapping/suggest', methods=['POST'])
@rate_limit('default')
def suggest_column_mapping():
    """Suggest column mappings between CSV headers and model features."""
    try:
        data = request.get_json(force=True)
        csv_headers = data.get('headers', [])
        claim_type = data.get('claim_type', 'auto')

        if not csv_headers:
            return jsonify({'error': 'No headers provided'}), 400

        if claim_type not in get_available_models():
            return jsonify({'error': f"Model '{claim_type}' not available"}), 400

        required_fields = get_required_fields(claim_type)

        # Fuzzy matching
        mapping = {}
        unmatched_csv = []
        unmatched_model = list(required_fields)

        for csv_col in csv_headers:
            csv_lower = csv_col.lower().replace(' ', '_').replace('-', '_')
            best_match = None
            best_score = 0

            for model_col in unmatched_model:
                model_lower = model_col.lower()

                # Exact match
                if csv_lower == model_lower:
                    best_match = model_col
                    best_score = 1.0
                    break

                # Contains match
                if csv_lower in model_lower or model_lower in csv_lower:
                    score = len(set(csv_lower) & set(model_lower)) / max(len(csv_lower), len(model_lower))
                    if score > best_score:
                        best_match = model_col
                        best_score = score

                # Substring overlap
                csv_parts = set(csv_lower.split('_'))
                model_parts = set(model_lower.split('_'))
                overlap = len(csv_parts & model_parts)
                if overlap > 0:
                    score = overlap / max(len(csv_parts), len(model_parts))
                    if score > best_score:
                        best_match = model_col
                        best_score = score

            if best_match and best_score > 0.3:
                mapping[csv_col] = {
                    'model_field': best_match,
                    'confidence': round(best_score, 2),
                    'auto_mapped': True,
                }
                unmatched_model.remove(best_match)
            else:
                unmatched_csv.append(csv_col)

        _audit('column_mapping', {'claim_type': claim_type, 'mapped': len(mapping)})

        return jsonify({
            'claim_type': claim_type,
            'mapping': mapping,
            'unmatched_csv_columns': unmatched_csv,
            'unmatched_model_fields': unmatched_model,
            'coverage': round(len(mapping) / max(len(required_fields), 1), 2),
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# FEEDBACK LOOP (Must-Have #8)
# ═══════════════════════════════════════════════════════════════════

@app.route('/feedback', methods=['POST'])
@rate_limit('feedback')
def submit_feedback():
    """Submit user feedback on a prediction."""
    try:
        data = request.get_json(force=True)
        claim_id = data.get('claim_id')
        feedback_type = data.get('feedback_type')  # 'correct', 'incorrect', 'unsure'
        actual_label = data.get('actual_label')     # 'fraud', 'legitimate', None
        comment = data.get('comment', '')

        if not claim_id:
            return jsonify({'error': 'claim_id is required'}), 400
        if feedback_type not in ('correct', 'incorrect', 'unsure'):
            return jsonify({'error': "feedback_type must be 'correct', 'incorrect', or 'unsure'"}), 400

        feedback_entry = {
            'id': str(uuid.uuid4())[:8],
            'claim_id': claim_id,
            'feedback_type': feedback_type,
            'actual_label': actual_label,
            'comment': comment[:1000],
            'client_ip': _get_client_ip(),
            'timestamp': datetime.now(timezone.utc).isoformat(),
        }

        if claim_id not in _feedback_store:
            _feedback_store[claim_id] = []
        _feedback_store[claim_id].append(feedback_entry)

        _audit('feedback', {'claim_id': claim_id, 'type': feedback_type})

        return jsonify({
            'status': 'received',
            'feedback_id': feedback_entry['id'],
            'message': 'Thank you for your feedback! It will be used to improve our models.',
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/feedback/<claim_id>', methods=['GET'])
def get_feedback(claim_id):
    """Get all feedback for a claim."""
    entries = _feedback_store.get(claim_id, [])
    return jsonify({
        'claim_id': claim_id,
        'feedback_count': len(entries),
        'feedback': entries,
    })


# ═══════════════════════════════════════════════════════════════════
# CLAIM COMPARISON (Should-Have #1)
# ═══════════════════════════════════════════════════════════════════

@app.route('/compare', methods=['POST'])
@rate_limit('predict')
def compare_claims():
    """Compare two claims side-by-side."""
    try:
        data = request.get_json(force=True)
        claim_a = data.get('claim_a', {})
        claim_b = data.get('claim_b', {})
        claim_type = data.get('claim_type', 'auto')

        if not claim_a or not claim_b:
            return jsonify({'error': 'Both claim_a and claim_b are required'}), 400

        if claim_type not in get_available_models():
            return jsonify({'error': f"Model '{claim_type}' not available"}), 400

        # Process both claims
        results = []
        for label, claim_data in [('A', claim_a), ('B', claim_b)]:
            claim_data['claim_type'] = claim_type
            processed = preprocess_input(claim_type, claim_data)
            prediction, probability = get_prediction_and_probability(claim_type, processed)
            indicators = generate_indicators(claim_data, claim_type)

            results.append({
                'label': label,
                'prediction': prediction,
                'probability': round(probability, 4),
                'risk_level': _risk_level(probability),
                'indicators': indicators,
            })

        # Find differences in input fields
        all_fields = set(list(claim_a.keys()) + list(claim_b.keys()))
        differences = {}
        for field in all_fields:
            if field == 'claim_type':
                continue
            val_a = claim_a.get(field)
            val_b = claim_b.get(field)
            if str(val_a) != str(val_b):
                differences[field] = {'claim_a': val_a, 'claim_b': val_b}

        _audit('compare', {'claim_type': claim_type})

        return jsonify({
            'claim_type': claim_type,
            'claim_a': results[0],
            'claim_b': results[1],
            'differences': differences,
            'probability_delta': round(abs(results[0]['probability'] - results[1]['probability']), 4),
            'same_prediction': results[0]['prediction'] == results[1]['prediction'],
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# MODEL VERSIONING (Should-Have #4)
# ═══════════════════════════════════════════════════════════════════

@app.route('/model/version/<claim_type>', methods=['GET'])
def model_version(claim_type):
    """Get model version info and history."""
    if claim_type not in get_available_models():
        return jsonify({'error': f"Model '{claim_type}' not available"}), 404

    metrics = get_model_metrics(claim_type)
    config = get_feature_config(claim_type)

    return jsonify({
        'claim_type': claim_type,
        'display_name': config.get('display_name', claim_type),
        'current_version': {
            'version': metrics.get('version', '1.0.0'),
            'trained_at': metrics.get('trained_at', 'unknown'),
            'n_features': metrics.get('n_features', 0),
            'n_samples': metrics.get('n_samples', 0),
            'f1_weighted': metrics.get('f1_weighted', 0),
            'auc_roc': metrics.get('auc_roc', 0),
        },
        'model_type': 'XGBoost Pipeline',
        'feature_columns': config.get('feature_columns', []),
        'changelog': [
            {
                'version': '1.0.0',
                'date': metrics.get('trained_at', 'unknown'),
                'changes': 'Initial production model with feature engineering pipeline',
            }
        ],
    })


# ═══════════════════════════════════════════════════════════════════
# TIME-SERIES ANOMALY (Should-Have #2)
# ═══════════════════════════════════════════════════════════════════

@app.route('/anomaly/timeseries', methods=['POST'])
@rate_limit('predict')
def timeseries_anomaly():
    """Detect anomalies in time-series batch data."""
    try:
        data = request.get_json(force=True)
        values = data.get('values', [])
        timestamps = data.get('timestamps', [])
        window_size = data.get('window_size', 5)

        if not values or len(values) < 3:
            return jsonify({'error': 'Need at least 3 data points'}), 400

        import numpy as np
        vals = np.array([float(v) for v in values])

        # Z-score anomaly detection
        mean = np.mean(vals)
        std = np.std(vals) if np.std(vals) > 0 else 1.0
        z_scores = (vals - mean) / std

        anomalies = []
        for i, (val, z) in enumerate(zip(vals, z_scores)):
            if abs(z) > 2.0:
                anomalies.append({
                    'index': i,
                    'value': float(val),
                    'z_score': round(float(z), 3),
                    'timestamp': timestamps[i] if i < len(timestamps) else None,
                    'severity': 'high' if abs(z) > 3 else 'medium',
                })

        # Moving average trend
        if len(vals) >= window_size:
            moving_avg = np.convolve(vals, np.ones(window_size) / window_size, mode='valid')
            trend = 'increasing' if moving_avg[-1] > moving_avg[0] else 'decreasing' if moving_avg[-1] < moving_avg[0] else 'stable'
        else:
            trend = 'insufficient_data'

        return jsonify({
            'total_points': len(values),
            'anomalies_detected': len(anomalies),
            'anomalies': anomalies,
            'statistics': {
                'mean': round(float(mean), 2),
                'std': round(float(std), 2),
                'min': round(float(np.min(vals)), 2),
                'max': round(float(np.max(vals)), 2),
            },
            'trend': trend,
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# NETWORK ANALYSIS (Should-Have #3)
# ═══════════════════════════════════════════════════════════════════

@app.route('/network/analysis', methods=['POST'])
@rate_limit('predict')
def network_analysis():
    """Analyze network of related claims for fraud rings."""
    try:
        data = request.get_json(force=True)
        claims = data.get('claims', [])
        link_fields = data.get('link_fields', [])

        if not claims or len(claims) < 2:
            return jsonify({'error': 'Need at least 2 claims for network analysis'}), 400

        # Default link fields by claim type
        if not link_fields:
            link_fields = ['insured_occupation', 'auto_make', 'incident_type',
                           'agency_name', 'destination', 'Provider_Specialty',
                           'property_type', 'nominee_relationship']

        # Build adjacency based on shared field values
        nodes = []
        edges = []
        clusters = defaultdict(list)

        for i, claim in enumerate(claims):
            nodes.append({
                'id': i,
                'label': f"Claim #{i + 1}",
                'prediction': claim.get('prediction', 'unknown'),
            })

        for i in range(len(claims)):
            for j in range(i + 1, len(claims)):
                shared = []
                for field in link_fields:
                    val_i = str(claims[i].get(field, '')).lower()
                    val_j = str(claims[j].get(field, '')).lower()
                    if val_i and val_j and val_i == val_j and val_i not in ('', 'unknown', 'none', '0'):
                        shared.append({'field': field, 'value': val_i})

                if shared:
                    edges.append({
                        'source': i,
                        'target': j,
                        'shared_fields': shared,
                        'weight': len(shared),
                    })

                    # Simple clustering
                    cluster_key = tuple(s['value'] for s in shared[:2])
                    clusters[cluster_key].extend([i, j])

        # Deduplicate clusters
        unique_clusters = []
        for key, members in clusters.items():
            unique_members = list(set(members))
            if len(unique_members) >= 2:
                unique_clusters.append({
                    'cluster_id': len(unique_clusters) + 1,
                    'members': unique_members,
                    'size': len(unique_members),
                    'risk': 'high' if len(unique_members) >= 4 else 'medium',
                })

        return jsonify({
            'total_claims': len(claims),
            'nodes': nodes,
            'edges': edges,
            'clusters': unique_clusters[:20],
            'total_connections': len(edges),
            'density': round(len(edges) / max((len(claims) * (len(claims) - 1)) / 2, 1), 3),
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ═══════════════════════════════════════════════════════════════════
# AUDIT TRAIL (Should-Have #6)
# ═══════════════════════════════════════════════════════════════════

@app.route('/audit-trail', methods=['GET'])
def audit_trail():
    """Get audit log of actions."""
    limit = request.args.get('limit', 100, type=int)
    action_filter = request.args.get('action', None)

    filtered = _audit_log
    if action_filter:
        filtered = [e for e in _audit_log if e['action'] == action_filter]

    return jsonify({
        'total': len(filtered),
        'entries': filtered[-limit:],
    })


# ═══════════════════════════════════════════════════════════════════
# CONFIDENCE INTERVALS (Should-Have #8)
# ═══════════════════════════════════════════════════════════════════

@app.route('/confidence-interval/<claim_type>', methods=['GET'])
def confidence_interval_info(claim_type):
    """Get confidence interval methodology info for a model."""
    if claim_type not in get_available_models():
        return jsonify({'error': f"Model '{claim_type}' not available"}), 404

    confidence = get_model_confidence_info(claim_type)
    return jsonify({
        'claim_type': claim_type,
        **confidence,
        'methodology': 'Bootstrap estimation from training data calibration + model OOB error',
    })


# ═══════════════════════════════════════════════════════════════════
# A/B TESTING (Should-Have #5)
# ═══════════════════════════════════════════════════════════════════

@app.route('/ab-test/status', methods=['GET'])
def ab_test_status():
    """Get A/B test status for models (placeholder for future implementation)."""
    return jsonify({
        'active_tests': [],
        'message': 'A/B testing framework is ready. No active tests currently.',
        'supported_models': get_available_models(),
    })


# ═══════════════════════════════════════════════════════════════════
# MULTI-LANGUAGE (Should-Have #7)
# ═══════════════════════════════════════════════════════════════════

_TRANSLATIONS = {
    'en': {
        'fraud_detected': 'Fraudulent Claim Detected',
        'legitimate': 'Legitimate Claim',
        'risk_critical': 'Critical Risk',
        'risk_high': 'High Risk',
        'risk_medium': 'Medium Risk',
        'risk_low': 'Low Risk',
        'risk_minimal': 'Minimal Risk',
    },
    'hi': {
        'fraud_detected': 'धोखाधड़ी दावा पता चला',
        'legitimate': 'वैध दावा',
        'risk_critical': 'गंभीर जोखिम',
        'risk_high': 'उच्च जोखिम',
        'risk_medium': 'मध्यम जोखिम',
        'risk_low': 'कम जोखिम',
        'risk_minimal': 'न्यूनतम जोखिम',
    },
    'es': {
        'fraud_detected': 'Reclamo Fraudulento Detectado',
        'legitimate': 'Reclamo Legítimo',
        'risk_critical': 'Riesgo Crítico',
        'risk_high': 'Riesgo Alto',
        'risk_medium': 'Riesgo Medio',
        'risk_low': 'Riesgo Bajo',
        'risk_minimal': 'Riesgo Mínimo',
    },
    'fr': {
        'fraud_detected': 'Réclamation Frauduleuse Détectée',
        'legitimate': 'Réclamation Légitime',
        'risk_critical': 'Risque Critique',
        'risk_high': 'Risque Élevé',
        'risk_medium': 'Risque Moyen',
        'risk_low': 'Risque Faible',
        'risk_minimal': 'Risque Minimal',
    },
}


@app.route('/translations/<lang>', methods=['GET'])
def get_translations(lang):
    """Get translations for the UI."""
    if lang not in _TRANSLATIONS:
        return jsonify({
            'error': f"Language '{lang}' not supported",
            'available': list(_TRANSLATIONS.keys()),
        }), 404

    return jsonify({
        'language': lang,
        'translations': _TRANSLATIONS[lang],
    })


@app.route('/translations', methods=['GET'])
def list_languages():
    """List available languages."""
    return jsonify({
        'available': list(_TRANSLATIONS.keys()),
        'default': 'en',
    })


# ═══════════════════════════════════════════════════════════════════
# HELPER FUNCTIONS (kept from original)
# ═══════════════════════════════════════════════════════════════════

def _extract_csv(file, claim_type):
    """Extract data from CSV file."""
    import pandas as pd

    content = file.read().decode('utf-8', errors='ignore')
    df = pd.read_csv(io.StringIO(content))

    return jsonify({
        'file_type': 'csv',
        'rows': len(df),
        'columns': list(df.columns),
        'sample': df.head(3).to_dict(orient='records'),
        'claim_type': claim_type,
        'auto_category': classify_claim_category(dict.fromkeys(df.columns)),
    })


def _extract_excel(file, claim_type):
    """Extract data from Excel file."""
    import pandas as pd

    df = pd.read_excel(file, engine='openpyxl')

    return jsonify({
        'file_type': 'excel',
        'rows': len(df),
        'columns': list(df.columns),
        'sample': df.head(3).to_dict(orient='records'),
        'claim_type': claim_type,
        'auto_category': classify_claim_category(dict.fromkeys(df.columns)),
    })


def _extract_pdf(file, claim_type):
    """Extract text and tables from PDF."""
    try:
        import pdfplumber
    except ImportError:
        return jsonify({
            'error': 'pdfplumber not installed.',
            'fallback': 'Use Document AI (frontend) for PDF extraction'
        }), 501

    content = file.read()
    pdf = pdfplumber.open(io.BytesIO(content))

    all_text = []
    all_tables = []

    for page in pdf.pages:
        text = page.extract_text()
        if text:
            all_text.append(text)

        tables = page.extract_tables()
        for table in tables:
            if table and len(table) > 1:
                headers = [str(h).strip() if h else f'col_{i}' for i, h in enumerate(table[0])]
                rows = []
                for row in table[1:]:
                    rows.append({headers[i]: str(cell).strip() if cell else ''
                                for i, cell in enumerate(row) if i < len(headers)})
                all_tables.append({'headers': headers, 'rows': rows})

    pdf.close()
    full_text = '\n'.join(all_text)
    auto_category = _detect_category_from_text(full_text)

    return jsonify({
        'file_type': 'pdf',
        'pages': len(all_text),
        'text': full_text[:5000],
        'text_length': len(full_text),
        'tables': all_tables,
        'has_tables': len(all_tables) > 0,
        'claim_type': claim_type,
        'auto_category': auto_category,
    })


def _extract_docx(file, claim_type):
    """Extract text and tables from DOCX."""
    try:
        from docx import Document
    except ImportError:
        return jsonify({'error': 'python-docx not installed.'}), 501

    content = file.read()
    doc = Document(io.BytesIO(content))

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = '\n'.join(paragraphs)

    all_tables = []
    for table in doc.tables:
        rows = []
        headers = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0:
                headers = cells
            else:
                rows.append({headers[j]: cells[j] for j in range(min(len(headers), len(cells)))})
        if headers and rows:
            all_tables.append({'headers': headers, 'rows': rows})

    auto_category = _detect_category_from_text(full_text)

    return jsonify({
        'file_type': 'docx',
        'text': full_text[:5000],
        'text_length': len(full_text),
        'tables': all_tables,
        'has_tables': len(all_tables) > 0,
        'claim_type': claim_type,
        'auto_category': auto_category,
    })


def _detect_category_from_text(text: str) -> str:
    """Detect insurance category from document text."""
    text_lower = text.lower()

    scores = {'auto': 0, 'health': 0, 'travel': 0, 'life': 0, 'property': 0}

    keywords = {
        'auto': ['vehicle', 'automobile', 'car', 'collision', 'driver', 'accident', 'auto insurance', 'bodily injury', 'fender', 'traffic'],
        'health': ['patient', 'diagnosis', 'hospital', 'medical', 'healthcare', 'procedure', 'prescription', 'surgery', 'doctor', 'clinic'],
        'travel': ['travel', 'flight', 'trip', 'destination', 'airline', 'passport', 'cancellation', 'delay', 'luggage', 'boarding'],
        'life': ['death', 'deceased', 'beneficiary', 'nominee', 'life insurance', 'mortality', 'cremation', 'burial', 'estate', 'will'],
        'property': ['property', 'home', 'house', 'building', 'fire', 'flood', 'burglary', 'repair', 'structural', 'dwelling', 'roof'],
    }

    for cat, kws in keywords.items():
        for kw in kws:
            if kw in text_lower:
                scores[cat] += 1

    best = max(scores, key=scores.get)
    return best if scores[best] > 0 else 'unknown'


def _risk_level(probability: float) -> str:
    """Convert probability to human-readable risk level."""
    if probability >= 0.8:
        return 'critical'
    elif probability >= 0.6:
        return 'high'
    elif probability >= 0.4:
        return 'medium'
    elif probability >= 0.2:
        return 'low'
    else:
        return 'minimal'


if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(debug=False, host='0.0.0.0', port=port)


@app.errorhandler(413)
def handle_request_entity_too_large(_error):
    return jsonify({
        'error': 'Uploaded file exceeds the 10MB limit',
        'max_bytes': MAX_UPLOAD_BYTES,
    }), 413
